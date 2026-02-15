import os
import time
from typing import Tuple, Dict
from docx import Document
from mutagen import File as MutagenFile

def read_docx(file_path: str) -> str:
    """Reads the content of a Word document."""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def clean_markdown(text: str) -> str:
    """Removes common markdown formatting like bold (**) and stars."""
    if not text:
        return ""
    # Remove bold markers
    text = text.replace("**", "").replace("__", "")
    # Remove bullet points that are just stars if they are at the start of lines
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if line.startswith('* '):
            line = "- " + line[2:]
        elif line.startswith('*'):
             line = line[1:].strip()
        lines.append(line)
    return '\n'.join(lines)

def save_docx(transcript: str, summary: str, output_path: str, title: str):
    """Saves transcript and summary to a Word document with simple formatting."""
    doc = Document()
    doc.add_heading(title, 0)
    
    if summary:
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(clean_markdown(summary))
        doc.add_page_break()
    
    if transcript:
        doc.add_heading('Transcript', level=1)
        for line in transcript.split('\n'):
            if line.strip():
                # Clean the line of any remaining markdown stars/bold
                doc.add_paragraph(clean_markdown(line))
    
    doc.save(output_path)

def save_json(content: str, output_path: str):
    """Saves content string (expecting JSON) to a file."""
    with open(output_path, 'w') as f:
        f.write(content)

def get_audio_duration(file_path: str) -> float:
    """Returns the duration of an audio file in seconds."""
    try:
        audio = MutagenFile(file_path)
        if audio is not None and audio.info is not None:
            return audio.info.length
        return 0.0
    except Exception:
        return 0.0

def save_assessment_docx(assessment_data: dict, output_path: str):
    """Saves agent assessment data (from JSON dict) to a Word document."""
    doc = Document()
    doc.add_heading('Agent Performance Assessment', 0)
    
    # 1. Summary
    doc.add_heading('Call Summary', level=1)
    doc.add_paragraph(assessment_data.get('call_summary', 'No summary provided.'))
    
    # 2. Performance Table
    doc.add_heading('Performance Evaluation', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Criterion'
    hdr_cells[1].text = 'Rating'
    
    performance = assessment_data.get('agent_performance', {})
    for criterion, rating in performance.items():
        row_cells = table.add_row().cells
        # Humanize criterion name
        criterion_name = criterion.replace('_', ' ').capitalize()
        row_cells[0].text = criterion_name
        row_cells[1].text = str(rating)
    
    # 3. Final Verdict
    doc.add_heading('Final Verdict', level=1)
    verdict = assessment_data.get('final_verdict', 'N/A')
    p = doc.add_paragraph()
    run = p.add_run(verdict)
    run.bold = True
    if verdict == 'Excellent':
        run.font.color.rgb = (0, 128, 0) # Green
    elif verdict == 'Poor':
        run.font.color.rgb = (255, 0, 0) # Red
        
    doc.save(output_path)

def format_timecode(seconds: float) -> str:
    """Formats seconds into [HH:MM:SS]."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"[{hours:02d}:{minutes:02d}:{secs:02d}]"
